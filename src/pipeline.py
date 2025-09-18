"""Command-line pipeline for isolating audio with ElevenLabs and transcribing Arabic speech.

This script expects the ELEVENLABS_API_KEY environment variable to be defined.
It will iterate over audio files in the provided input directory, call the
ElevenLabs audio isolation endpoint, transcribe the cleaned audio with a local
Whisper model, and emit Word documents alongside plain-text transcripts.
"""
from __future__ import annotations

import argparse
import logging
import os
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator, List, Sequence, Tuple

from docx import Document
from docx.shared import Pt
from elevenlabs.client import ElevenLabs
from elevenlabs.core import ApiError
from faster_whisper import WhisperModel

LOGGER = logging.getLogger("pipeline")

SUPPORTED_EXTENSIONS = {".mp3", ".wav", ".m4a", ".flac", ".ogg", ".webm"}


@dataclass
class PipelineConfig:
    """Runtime configuration for the isolation and transcription pipeline."""

    input_dir: Path
    output_dir: Path
    language: str = "ar"
    model_size: str = "small"
    device: str = "cpu"
    compute_type: str = "int8"
    chunk_size: int = 1024 * 64
    skip_existing: bool = True
    model_dir: Path | None = None

    def ensure_output_subdirs(self) -> Tuple[Path, Path, Path]:
        """Create the output sub-directories used by the pipeline."""

        isolated_dir = self.output_dir / "isolated_audio"
        text_dir = self.output_dir / "transcripts"
        docx_dir = self.output_dir / "word_docs"
        for directory in (self.output_dir, isolated_dir, text_dir, docx_dir):
            directory.mkdir(parents=True, exist_ok=True)
        return isolated_dir, text_dir, docx_dir


class ElevenLabsVoiceIsolator:
    """Thin wrapper around the ElevenLabs audio isolation endpoint."""

    def __init__(self, api_key: str, chunk_size: int = 1024 * 64) -> None:
        self._client = ElevenLabs(api_key=api_key)
        self._chunk_size = chunk_size

    def isolate(self, source: Path, target: Path) -> None:
        """Run the ElevenLabs isolator on ``source`` and write to ``target``.

        Parameters
        ----------
        source:
            Path to the noisy input audio file.
        target:
            Destination path where the cleaned audio will be written.
        """

        target.parent.mkdir(parents=True, exist_ok=True)
        start = time.time()
        LOGGER.info("Calling ElevenLabs isolator for %s", source.name)
        try:
            with open(source, "rb") as payload, open(target, "wb") as sink:
                stream = self._client.audio_isolation.convert(
                    audio=(source.name, payload, "audio/mpeg"),
                    file_format="other",
                    request_options={"chunk_size": self._chunk_size},
                )
                for chunk in stream:
                    sink.write(chunk)
        except ApiError as exc:  # pragma: no cover - network failure path
            LOGGER.error("ElevenLabs API error for %s: %s", source.name, exc)
            raise
        except Exception:
            LOGGER.exception("Unexpected error while isolating %s", source)
            raise
        elapsed = time.time() - start
        size_kb = target.stat().st_size / 1024 if target.exists() else 0
        LOGGER.info(
            "Finished isolation for %s in %.1fs (%.1f KiB)",
            source.name,
            elapsed,
            size_kb,
        )


class WhisperTranscriber:
    """Wrapper around ``faster-whisper`` to transcribe Arabic speech."""

    def __init__(
        self,
        model_size: str,
        device: str,
        compute_type: str,
        model_dir: Path | None = None,
    ) -> None:
        LOGGER.info(
            "Loading Whisper model '%s' on %s (%s)", model_size, device, compute_type
        )
        kwargs = {}
        if model_dir is not None:
            kwargs["download_root"] = str(model_dir)
        self._model = WhisperModel(
            model_size,
            device=device,
            compute_type=compute_type,
            **kwargs,
        )

    def transcribe(self, audio_path: Path, language: str) -> "TranscriptionOutcome":
        segments: List[SegmentInfo] = []
        text_parts: List[str] = []
        start = time.time()
        LOGGER.info("Transcribing %s", audio_path.name)
        raw_segments, info = self._model.transcribe(
            str(audio_path),
            language=language,
            task="transcribe",
            beam_size=5,
            vad_filter=True,
        )
        for segment in raw_segments:
            seg = SegmentInfo(
                start=segment.start,
                end=segment.end,
                text=segment.text.strip(),
            )
            if seg.text:
                segments.append(seg)
                text_parts.append(seg.text)
        elapsed = time.time() - start
        detected_language = getattr(info, "language", None)
        duration = getattr(info, "duration", None)
        duration_desc = f"{duration:.1f}s" if isinstance(duration, (int, float)) else "unknown"
        LOGGER.info(
            "Finished transcription for %s in %.1fs (%d segments, detected language: %s, audio duration: %s)",
            audio_path.name,
            elapsed,
            len(segments),
            detected_language or "unknown",
            duration_desc,
        )
        return TranscriptionOutcome(
            text_parts=text_parts,
            segments=segments,
            detected_language=detected_language,
            duration=duration,
        )


@dataclass
class SegmentInfo:
    start: float
    end: float
    text: str


@dataclass
class TranscriptionOutcome:
    text_parts: List[str]
    segments: List[SegmentInfo]
    detected_language: str | None
    duration: float | None


class TranscriptWriter:
    """Persist transcripts in both plain text and Word document formats."""

    def __init__(self, text_dir: Path, docx_dir: Path) -> None:
        self._text_dir = text_dir
        self._docx_dir = docx_dir

    def write(
        self,
        *,
        source_file: Path,
        isolated_file: Path,
        language: str,
        model_size: str,
        outcome: TranscriptionOutcome,
    ) -> Tuple[Path, Path]:
        text_output = self._text_dir / f"{source_file.stem}.txt"
        doc_output = self._docx_dir / f"{source_file.stem}.docx"
        text_output.write_text("\n".join(outcome.text_parts), encoding="utf-8")
        LOGGER.info("Saved text transcript to %s", text_output)

        document = Document()
        document.add_heading(f"Transcript for {source_file.name}", level=1)
        metadata = document.add_paragraph()
        metadata.add_run("Original file: ").bold = True
        metadata.add_run(source_file.name)
        metadata.add_run("\nIsolated audio: ").bold = True
        metadata.add_run(isolated_file.name)
        metadata.add_run("\nLanguage: ").bold = True
        metadata.add_run(language)
        metadata.add_run("\nModel: ").bold = True
        metadata.add_run(model_size)
        if outcome.detected_language:
            metadata.add_run("\nDetected language: ").bold = True
            metadata.add_run(outcome.detected_language)
        if isinstance(outcome.duration, (int, float)):
            metadata.add_run("\nAudio duration: ").bold = True
            metadata.add_run(format_timestamp(outcome.duration))

        document.add_paragraph("")
        document.add_heading("Segments", level=2)
        for index, segment in enumerate(outcome.segments, start=1):
            paragraph = document.add_paragraph(style="List Number")
            run = paragraph.add_run(
                f"[{format_timestamp(segment.start)} - {format_timestamp(segment.end)}] {segment.text}"
            )
            run.font.size = Pt(11)
        document.save(doc_output)
        LOGGER.info("Saved Word transcript to %s", doc_output)
        return text_output, doc_output


def format_timestamp(value: float) -> str:
    hours, remainder = divmod(int(value), 3600)
    minutes, seconds = divmod(remainder, 60)
    millis = int(round((value - int(value)) * 1000))
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}.{millis:03d}"
    return f"{minutes:02d}:{seconds:02d}.{millis:03d}"


def iter_audio_files(input_dir: Path) -> Iterator[Path]:
    for path in sorted(input_dir.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield path


def configure_logging(verbose: bool = False) -> None:
    logging.basicConfig(
        level=logging.DEBUG if verbose else logging.INFO,
        format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
    )


def run_pipeline(config: PipelineConfig, *, verbose: bool = False) -> None:
    configure_logging(verbose=verbose)

    api_key = os.getenv("ELEVENLABS_API_KEY")
    if not api_key:
        LOGGER.error("ELEVENLABS_API_KEY environment variable is not set.")
        sys.exit(1)

    isolated_dir, text_dir, docx_dir = config.ensure_output_subdirs()
    isolator = ElevenLabsVoiceIsolator(api_key, chunk_size=config.chunk_size)
    transcriber = WhisperTranscriber(
        config.model_size,
        config.device,
        config.compute_type,
        model_dir=config.model_dir,
    )
    writer = TranscriptWriter(text_dir, docx_dir)

    audio_files = list(iter_audio_files(config.input_dir))
    if not audio_files:
        LOGGER.warning("No audio files found in %s", config.input_dir)
        return

    for audio_path in audio_files:
        LOGGER.info("Processing %s", audio_path)
        isolated_path = isolated_dir / f"{audio_path.stem}_isolated.wav"
        if not (config.skip_existing and isolated_path.exists() and isolated_path.stat().st_size > 0):
            isolator.isolate(audio_path, isolated_path)
        else:
            LOGGER.info("Skipping isolation for %s (existing file)", audio_path.name)
        outcome = transcriber.transcribe(isolated_path, config.language)
        writer.write(
            source_file=audio_path,
            isolated_file=isolated_path,
            language=config.language,
            model_size=config.model_size,
            outcome=outcome,
        )


def parse_args(argv: Sequence[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Isolate audio with ElevenLabs and transcribe Arabic speech to Word documents.",
    )
    parser.add_argument(
        "--input-dir",
        type=Path,
        default=Path("."),
        help="Directory containing input audio files.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("outputs"),
        help="Directory where isolated audio and transcripts will be stored.",
    )
    parser.add_argument(
        "--language",
        default="ar",
        help="Language code for transcription (default: ar).",
    )
    parser.add_argument(
        "--model-size",
        default="small",
        help="Whisper model size to load (tiny, base, small, medium, large-v2, ...).",
    )
    parser.add_argument(
        "--device",
        default="cpu",
        help="Device to use for Whisper inference (cpu, cuda).",
    )
    parser.add_argument(
        "--compute-type",
        default="int8",
        help="Quantization for Whisper inference (int8, int8_float16, float16, float32).",
    )
    parser.add_argument(
        "--chunk-size",
        type=int,
        default=1024 * 64,
        help="Chunk size (in bytes) to use when downloading isolated audio.",
    )
    parser.add_argument(
        "--no-skip-existing",
        action="store_true",
        help="Re-run isolation even if cleaned audio already exists.",
    )
    parser.add_argument(
        "--model-dir",
        type=Path,
        default=None,
        help="Optional directory to cache Whisper models.",
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Enable debug logging.",
    )
    return parser.parse_args(argv)


def main(argv: Sequence[str] | None = None) -> None:
    args = parse_args(argv or sys.argv[1:])
    config = PipelineConfig(
        input_dir=args.input_dir,
        output_dir=args.output_dir,
        language=args.language,
        model_size=args.model_size,
        device=args.device,
        compute_type=args.compute_type,
        chunk_size=args.chunk_size,
        skip_existing=not args.no_skip_existing,
        model_dir=args.model_dir,
    )
    run_pipeline(config, verbose=args.verbose)


if __name__ == "__main__":
    main()
